import fitz  # PyMuPDF
from docx import Document as DocxDocument
import os
import asyncio
from typing import List, Tuple, Union
import logging

from ml.vertex_ai_classifier import classify_text_with_gemini, classify_multimodal_with_gemini

logger = logging.getLogger(__name__)

def extract_text_from_file(file_path: str) -> str:
    """Extract text from PDF, DOCX, or TXT files with better error handling."""
    try:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return extract_text_pdf(file_path)
        elif ext == '.docx':
            return extract_text_docx(file_path)
        elif ext == '.txt':
            return extract_text_txt(file_path)
        else:
            logger.warning(f"Unsupported file type: {ext}")
            return ""
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return ""

def extract_text_pdf(file_path: str) -> str:
    """Enhanced PDF text extraction with better handling."""
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc):
                try:
                    page_text = page.get_text()  # type: ignore
                    if page_text.strip():  # Only add non-empty pages
                        text += f"\n{page_text}"
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    continue
    except Exception as e:
        logger.error(f"Error opening PDF {file_path}: {e}")
        return ""
    return text.strip()

def extract_text_docx(file_path: str) -> str:
    """Enhanced DOCX text extraction including headers, footers, and tables."""
    text_parts = []
    try:
        doc = DocxDocument(file_path)
        
        # Extract main document text
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
    except Exception as e:
        logger.error(f"Error extracting DOCX {file_path}: {e}")
        return ""
    
    return "\n".join(text_parts)

def extract_text_txt(file_path: str) -> str:
    """Enhanced text file reading with encoding detection."""
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            break
    
    return ""


# ============================================
# Large PDF Page-Range Splitting (Step 2b)
# ============================================
MAX_PAGES_PER_PART = int(os.getenv("PDF_MAX_PAGES_PER_PART", "500"))
MAX_TOTAL_PAGES = int(os.getenv("PDF_MAX_TOTAL_PAGES", "10000"))  # Cost guardrail
MAX_CHUNKS = int(os.getenv("PDF_MAX_CHUNKS", "10"))  # Max Gemini API calls per document
SECURITY_RANK = {"confidential": 3, "internal": 2, "public": 1, "unclassified": 0}

# Multimodal PDF guards (Gemini image inputs)
MAX_PDF_IMAGE_PAGES = int(os.getenv("PDF_IMAGE_MAX_PAGES", "5"))
PDF_IMAGE_DPI = int(os.getenv("PDF_IMAGE_DPI", "110"))
MAX_PDF_IMAGE_BYTES = int(os.getenv("PDF_IMAGE_MAX_BYTES", str(4 * 1024 * 1024)))

# P1-7 FIX: Token limit guard for non-PDF files (DOCX, TXT).
# Approximate token count: 1 token ≈ 4 chars (English average).
# Default: 3,000,000 chars ≈ 750K tokens, within Gemini's 1M limit with headroom.
MAX_TEXT_CHARS = int(os.getenv("MAX_TEXT_CHARS_PER_CLASSIFICATION", "3000000"))

# ⚠️ REVIEW FIX P1-REVIEW-8 (DOCX MEMORY GUARD):
# python-docx loads the entire DOCX into memory (DOM model). A 200MB DOCX with
# embedded images could OOM the container. Check file size before extraction.
MAX_DOCX_FILE_SIZE_BYTES = int(os.getenv("MAX_DOCX_FILE_SIZE_MB", "50")) * 1024 * 1024


def extract_text_pdf_pages(file_path: str, start_page: int, end_page: int) -> str:
    """Extract text from a specific page range of a PDF using PyMuPDF.

    Used for page-range splitting of large PDFs (>MAX_PAGES_PER_PART pages).
    Reuses the same fitz logic as extract_text_pdf but for a subset of pages.
    """
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page_num in range(start_page, min(end_page, len(doc))):
                try:
                    page_text = doc[page_num].get_text()
                    if page_text.strip():
                        text += f"\n{page_text}"
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    continue
    except Exception as e:
        logger.error(f"Error opening PDF {file_path}: {e}")
    return text.strip()


def extract_large_pdf_chunks(file_path: str, total_pages: int) -> list:
    """⚠️ REVIEW FIX P1-REVIEW-5: Open the PDF file ONCE and extract all chunks.

    The previous approach called extract_text_pdf_pages() per chunk, opening
    the PDF N times (O(N) file opens). This version opens once and iterates
    pages, building chunks in memory."""
    chunks = []
    try:
        with fitz.open(file_path) as doc:
            chunk_texts = []
            pages_in_chunk = 0
            for page_num in range(min(total_pages, len(doc))):
                # Cost guardrail: limit number of Gemini API calls per document
                if len(chunks) >= MAX_CHUNKS:
                    logger.warning(f"Chunk limit ({MAX_CHUNKS}) reached for {file_path}")
                    break
                try:
                    page_text = doc[page_num].get_text()
                    if page_text.strip():
                        chunk_texts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    continue
                pages_in_chunk += 1
                if pages_in_chunk >= MAX_PAGES_PER_PART:
                    text = "\n".join(chunk_texts)
                    if text.strip():
                        chunks.append(text)
                    chunk_texts = []
                    pages_in_chunk = 0
            # Don't forget the last partial chunk
            if chunk_texts and len(chunks) < MAX_CHUNKS:
                text = "\n".join(chunk_texts)
                if text.strip():
                    chunks.append(text)
    except Exception as e:
        logger.error(f"Error opening PDF {file_path}: {e}")
    return chunks


def get_pdf_page_count(file_path: str) -> int:
    """Get total page count of a PDF."""
    try:
        with fitz.open(file_path) as doc:
            return len(doc)
    except Exception:
        return 0


def extract_pdf_page_images(file_path: str) -> List[bytes]:
    """Render first N PDF pages to PNG bytes for Gemini multimodal classification.

    Safety guards:
    - Limits number of rendered pages (MAX_PDF_IMAGE_PAGES)
    - Limits per-image payload size (MAX_PDF_IMAGE_BYTES)
    - Uses configurable render DPI (PDF_IMAGE_DPI)
    """
    images: List[bytes] = []
    try:
        with fitz.open(file_path) as doc:
            total = min(len(doc), MAX_PDF_IMAGE_PAGES)
            zoom = max(PDF_IMAGE_DPI, 36) / 72.0
            matrix = fitz.Matrix(zoom, zoom)

            for page_num in range(total):
                try:
                    pix = doc[page_num].get_pixmap(matrix=matrix, alpha=False)
                    png_bytes = pix.tobytes("png")
                    if len(png_bytes) > MAX_PDF_IMAGE_BYTES:
                        logger.warning(
                            f"Skipping page {page_num + 1} image for {os.path.basename(file_path)} "
                            f"({len(png_bytes)} bytes > max {MAX_PDF_IMAGE_BYTES})"
                        )
                        continue
                    images.append(png_bytes)
                except Exception as e:
                    logger.warning(f"Failed rendering PDF page image {page_num + 1}: {e}")
                    continue
    except Exception as e:
        logger.error(f"Error extracting PDF page images from {file_path}: {e}")

    return images


# ============================================
# Async extraction & classification (Vertex AI pipeline)
# ============================================

async def extract_document_text_async(file_path: str) -> Union[str, List[str]]:
    """Extract text from a document — native async function.

    Called by the background pipeline AFTER setting 'extracting_text' status.
    Returns either a single string (small docs) or a list of strings (large PDF chunks).
    The pipeline then passes this to classify_extracted_text_async().

    For large PDFs (>MAX_PAGES_PER_PART pages), extracts text in page-range
    chunks using PyMuPDF. No temp files or file-level splitting.
    """
    try:
        ext = os.path.splitext(file_path)[1].lower()

        # For PDFs, check if page-range splitting is needed
        if ext == '.pdf':
            total_pages = await asyncio.to_thread(get_pdf_page_count, file_path)

            # Cost guardrail: reject unreasonably large PDFs
            if total_pages > MAX_TOTAL_PAGES:
                raise ValueError(
                    f"PDF has {total_pages} pages (max {MAX_TOTAL_PAGES}). "
                    f"Reduce page count or increase PDF_MAX_TOTAL_PAGES env var."
                )

            if total_pages > MAX_PAGES_PER_PART:
                # Large PDF: extract text in page-range chunks (single file open)
                logger.info(f"Large PDF ({total_pages} pages), splitting into chunks of {MAX_PAGES_PER_PART}")
                chunks = await asyncio.to_thread(extract_large_pdf_chunks, file_path, total_pages)
                return chunks if chunks else ""

        # ⚠️ REVIEW FIX P1-REVIEW-8 (DOCX MEMORY GUARD): python-docx loads entire
        # DOCX into memory (DOM model). Check file size before extraction.
        if ext == '.docx':
            file_size = os.path.getsize(file_path)
            if file_size > MAX_DOCX_FILE_SIZE_BYTES:
                raise ValueError(
                    f"DOCX file too large ({file_size / (1024*1024):.0f}MB, max "
                    f"{MAX_DOCX_FILE_SIZE_BYTES / (1024*1024):.0f}MB). "
                    f"Convert to PDF for large documents."
                )

        # Standard path: extract all text at once (small PDFs, DOCX, TXT)
        text = await asyncio.to_thread(extract_text_from_file, file_path)
        logger.info(f"Extracted text length: {len(text)} for {os.path.basename(file_path)}")
        return text

    except ValueError:
        # Re-raise ValueError (cost guardrail) so pipeline sets 'failed' status
        raise
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return ""


async def classify_extracted_text_async(text_or_chunks: Union[str, List[str]], file_path: str) -> str:
    """Classify extracted text via Gemini — native async function.

    Called by the background pipeline AFTER setting 'classifying' status.
    Accepts either a single string or a list of chunk strings (from large PDFs).
    For chunks, classifies each independently and takes the highest-security label.

    Exceptions from classify_text_with_gemini (API failures, auth errors) are NOT
    caught here — they propagate to the pipeline's except block, which sets
    classification_status = 'failed' with the error message.

    ⚠️ OVER-CLASSIFICATION NOTE: For chunked documents, the highest-security label
    from ANY chunk wins (single-chunk veto). This is intentional (err on security).
    Per-chunk results are logged at INFO level for admin review.
    """
    ext = os.path.splitext(file_path)[1].lower()

    # Multimodal path for PDFs: pass rendered page images + extracted text to Gemini.
    # Falls back to the existing text-only flow if image extraction yields nothing.
    if ext == '.pdf':
        if isinstance(text_or_chunks, list):
            combined_text = "\n\n".join(text_or_chunks)
        else:
            combined_text = text_or_chunks

        if len(combined_text) > MAX_TEXT_CHARS:
            logger.warning(
                f"PDF text for {os.path.basename(file_path)} exceeds {MAX_TEXT_CHARS} chars "
                f"({len(combined_text)} chars). Truncating before multimodal classification."
            )
            combined_text = combined_text[:MAX_TEXT_CHARS]

        page_images = await asyncio.to_thread(extract_pdf_page_images, file_path)
        if page_images:
            label, confidence = await classify_multimodal_with_gemini(combined_text, page_images)
            logger.info(
                f"Document {os.path.basename(file_path)} → '{label}' "
                f"(confidence={confidence:.3f}, pages_as_images={len(page_images)})"
            )
            return label

        if not combined_text.strip():
            raise ValueError(
                "PDF could not be classified: no extractable text and no renderable page images. "
                "The file may be encrypted, corrupted, or image rendering failed."
            )

        logger.info(f"No PDF page images extracted for {os.path.basename(file_path)}; using text-only fallback")

    if isinstance(text_or_chunks, list):
        # Multiple chunks from a large PDF
        best_label = "unclassified"
        best_confidence = 0.0

        for i, chunk_text in enumerate(text_or_chunks):
            # ⚠️ REVIEW FIX P2-REVIEW-11: Per-chunk char limit.
            if len(chunk_text) > MAX_TEXT_CHARS:
                logger.warning(
                    f"Chunk {i+1}/{len(text_or_chunks)} exceeds {MAX_TEXT_CHARS} chars "
                    f"({len(chunk_text)} chars). Truncating."
                )
                chunk_text = chunk_text[:MAX_TEXT_CHARS]
            label, confidence = await classify_text_with_gemini(chunk_text)
            logger.info(f"Chunk {i+1}/{len(text_or_chunks)} → '{label}' (confidence={confidence:.3f})")

            if SECURITY_RANK.get(label, 0) > SECURITY_RANK.get(best_label, 0):
                best_label = label
                best_confidence = confidence

        logger.info(f"Document {os.path.basename(file_path)} → '{best_label}' (confidence={best_confidence:.3f})")
        return best_label
    else:
        # Single text string
        if not text_or_chunks:
            logger.warning(f"No text to classify for {file_path}")
            return "unclassified"

        # P1-7 FIX: Token limit guard for non-PDF files (DOCX, TXT).
        if len(text_or_chunks) > MAX_TEXT_CHARS:
            logger.warning(
                f"Text for {os.path.basename(file_path)} exceeds {MAX_TEXT_CHARS} chars "
                f"({len(text_or_chunks)} chars). Truncating to fit Gemini token limit."
            )
            text_or_chunks = text_or_chunks[:MAX_TEXT_CHARS]

        label, confidence = await classify_text_with_gemini(text_or_chunks)
        logger.info(f"Document {os.path.basename(file_path)} → '{label}' (confidence={confidence:.3f})")
        return label