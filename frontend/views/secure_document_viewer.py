"""
Secure Document Viewer with YOLOv8 monitoring for confidential documents.
"""
from PyQt6.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, QPushButton, 
                             QTextEdit, QLabel, QMessageBox, QWidget, QScrollArea)
from PyQt6.QtCore import Qt, QTimer
from PyQt6.QtGui import QPixmap, QImage, QPainter, QFont, QColor
import qtawesome as qta
import tempfile
import os
import sys
import platform
from pathlib import Path
from datetime import datetime

from workers.detection_worker import DetectionWorker
from widgets.security_overlay import SecurityBlockOverlay, SecurityStatusBar, WatermarkOverlay

# Windows-specific imports for screen capture protection
if platform.system() == 'Windows':
    try:
        import ctypes
        from ctypes import wintypes
        WDA_EXCLUDEFROMCAPTURE = 0x00000011
        WINDOWS_CAPTURE_PROTECTION_AVAILABLE = True
    except ImportError:
        WINDOWS_CAPTURE_PROTECTION_AVAILABLE = False
else:
    WINDOWS_CAPTURE_PROTECTION_AVAILABLE = False


class SecureDocumentViewer(QDialog):
    """
    Secure document viewer with YOLOv8 monitoring for confidential documents.
    Monitors for person presence and cell phone detection.
    """
    
    def __init__(self, document_data: dict, file_content: bytes, api_client, parent=None):
        super().__init__(parent)
        self.document_data = document_data
        self.file_content = file_content
        self.filename = document_data.get('filename', 'Document')
        self.file_extension = Path(self.filename).suffix.lower()
        self.classification = document_data.get('classification', 'unclassified')
        self.api_client = api_client
        
        # Security monitoring
        self.detection_worker = None
        self.is_monitoring = False
        self.person_present = False
        self.phone_detection_logged = False
        self.model_initialized = False
        
        self.setWindowTitle(f"🔒 Secure View - {self.filename}")
        self.setMinimumSize(900, 700)
        
        # Set window flags (same for all documents - WindowStaysOnTopHint breaks UI)
        self.setWindowFlags(
            Qt.WindowType.Window |
            Qt.WindowType.WindowMaximizeButtonHint |
            Qt.WindowType.WindowMinimizeButtonHint |
            Qt.WindowType.WindowCloseButtonHint
        )
        
        # Additional security attributes for confidential documents
        if self.classification.lower() == 'confidential':
            # Install event filter to catch all key events including system keys
            self.installEventFilter(self)
            
            # Apply Windows screen capture protection
            self._apply_screen_capture_protection()
        
        self.setup_ui()
        
        # For confidential documents, wait for security initialization before loading content
        if self.classification.lower() == 'confidential':
            self._show_security_initialization_message()
        else:
            self.load_document()
        
        # Start monitoring if confidential
        if self.classification.lower() == 'confidential':
            self.start_monitoring()
        
        # Open maximized
        self.showMaximized()
    
    def setup_ui(self):
        main_layout = QVBoxLayout(self)
        
        # Header with document info
        header_layout = QHBoxLayout()
        
        info_label = QLabel(f"<b>Document:</b> {self.filename}")
        info_label.setStyleSheet("font-size: 14px; padding: 10px;")
        header_layout.addWidget(info_label)
        
        classification_label = QLabel(f"<b>Classification:</b> {self.classification.upper()}")
        classification_label.setStyleSheet(f"""
            font-size: 14px; 
            padding: 10px; 
            background-color: {self._get_classification_color(self.classification)};
            border-radius: 5px;
            color: white;
            font-weight: bold;
        """)
        header_layout.addWidget(classification_label)
        
        header_layout.addStretch()
        
        close_button = QPushButton(qta.icon('fa5s.times', color='red'), "Close")
        close_button.clicked.connect(self.close)
        header_layout.addWidget(close_button)
        
        main_layout.addLayout(header_layout)
        
        # Security status bar (only for confidential)
        if self.classification.lower() == 'confidential':
            self.status_bar = SecurityStatusBar(self)
            main_layout.addWidget(self.status_bar)
        
        # Content area
        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        main_layout.addWidget(self.content_widget)
        
        # Warning label
        if self.classification.lower() == 'confidential':
            warning_label = QLabel("🔒 CONFIDENTIAL: This document is under active security monitoring")
        else:
            warning_label = QLabel("⚠️ This document is displayed in read-only mode for security purposes")
            
        warning_label.setStyleSheet("""
            background-color: #fff3cd;
            color: #856404;
            padding: 10px;
            border: 1px solid #ffc107;
            border-radius: 5px;
            font-size: 12px;
        """)
        warning_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(warning_label)
        
        # Security overlay (only for confidential)
        if self.classification.lower() == 'confidential':
            self.security_overlay = SecurityBlockOverlay(self)
            
            # Add dynamic watermark overlay
            self.watermark_overlay = WatermarkOverlay(self, self.api_client)
            self.watermark_overlay.setGeometry(self.rect())
    
    def _show_security_initialization_message(self):
        """Show waiting message while security system initializes."""
        waiting_label = QLabel("""
            <div style='text-align: center; padding: 50px;'>
                <h2>🔒 Security System Initializing</h2>
                <p>Loading YOLOv8 object detection model...</p>
                <p>Please wait while we prepare secure viewing.</p>
                <div style='margin: 20px 0;'>
                    <span style='font-size: 24px;'>⏳</span>
                </div>
            </div>
        """)
        waiting_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        waiting_label.setStyleSheet("""
            QLabel {
                background-color: #f8f9fa;
                border: 2px solid #007bff;
                border-radius: 10px;
                color: #495057;
                font-size: 14px;
            }
        """)
        self.content_layout.addWidget(waiting_label)
        self.waiting_label = waiting_label
    
    def start_monitoring(self):
        """Start YOLOv8 monitoring."""
        try:
            self.status_bar.update_status("Initializing security monitoring...")
            
            # Create and configure detection worker
            self.detection_worker = DetectionWorker(camera_index=0)
            
            # Connect signals
            self.detection_worker.person_detected.connect(self.on_person_detection)
            self.detection_worker.phone_detected.connect(self.on_phone_detection)
            self.detection_worker.camera_error.connect(self.on_camera_error)
            self.detection_worker.detection_status.connect(self.on_detection_status)
            self.detection_worker.model_initialized.connect(self.on_model_initialized)
            
            # Start the worker thread
            self.detection_worker.start()
            self.is_monitoring = True
            
        except Exception as e:
            QMessageBox.critical(self, "Monitoring Error", 
                               f"Failed to start security monitoring: {str(e)}")
            self.close()
    
    def on_person_detection(self, person_present: bool):
        """Handle person detection status."""
        self.person_present = person_present
        
        if person_present:
            # Person detected - hide block
            self.security_overlay.hide_block()
        else:
            # No person - show block and log
            self.security_overlay.show_person_missing_block()
            
            # Log the event
            self.api_client.log_security_event(
                activity_type="no_person_detected",
                metadata={
                    "document_id": self.document_data.get('id'),
                    "document_name": self.filename,
                    "classification": self.classification
                }
            )
    
    def on_phone_detection(self, detection_data: dict):
        """Handle phone detection - keeps blocking continuously while phone is detected."""
        # Show block overlay (will stay visible)
        self.security_overlay.show_phone_detected_block()
        
        # Log to backend (only once per detection session)
        if not self.phone_detection_logged:
            self.api_client.log_security_event(
                activity_type="phone_detected",
                metadata={
                    "document_id": self.document_data.get('id'),
                    "document_name": self.filename,
                    "classification": self.classification,
                    "detection_timestamp": detection_data.get('timestamp'),
                    "confidence": detection_data.get('confidence')
                }
            )
            self.phone_detection_logged = True
        
        # Note: Block stays visible continuously until phone is removed
        # Detection worker will stop emitting phone_detected when phone is gone
        # Then we can check after a delay if phone is truly removed
        if hasattr(self, '_phone_check_timer'):
            self._phone_check_timer.stop()
        
        self._phone_check_timer = QTimer()
        self._phone_check_timer.timeout.connect(self._check_phone_removed)
        self._phone_check_timer.setSingleShot(True)
        self._phone_check_timer.start(3000)  # Check after 3 seconds of no new detections
    
    def _check_phone_removed(self):
        """Check if phone has been removed (called when no new phone detections)."""
        # Hide block and reset for new detections
        if self.person_present:
            self.security_overlay.hide_block()
            self.phone_detection_logged = False  # Allow new detection sessions
    
    def on_camera_error(self, error_message: str):
        """Handle camera errors."""
        QMessageBox.critical(self, "Camera Error", error_message)
        self.status_bar.update_status(f"Error: {error_message}")
    
    def on_detection_status(self, status: str):
        """Update status bar with detection status."""
        if hasattr(self, 'status_bar'):
            self.status_bar.update_status(status)
    
    def on_model_initialized(self, success: bool):
        """Handle model initialization completion."""
        self.model_initialized = success
        
        if success:
            # Remove waiting message and load document
            if hasattr(self, 'waiting_label'):
                self.content_layout.removeWidget(self.waiting_label)
                self.waiting_label.deleteLater()
                delattr(self, 'waiting_label')
            
            # Now load and display the document
            self.load_document()
        else:
            # Model initialization failed - show error
            if hasattr(self, 'waiting_label'):
                self.waiting_label.setText("""
                    <div style='text-align: center; padding: 50px;'>
                        <h2>❌ Security System Error</h2>
                        <p>Failed to initialize YOLOv8 model.</p>
                        <p>Please check your installation and try again.</p>
                        <div style='margin: 20px 0;'>
                            <span style='font-size: 24px;'>⚠️</span>
                        </div>
                    </div>
                """)
                self.waiting_label.setStyleSheet("""
                    QLabel {
                        background-color: #f8d7da;
                        border: 2px solid #dc3545;
                        border-radius: 10px;
                        color: #721c24;
                        font-size: 14px;
                    }
                """)
    
    def _apply_screen_capture_protection(self):
        """Apply Windows screen capture protection using SetWindowDisplayAffinity."""
        if not WINDOWS_CAPTURE_PROTECTION_AVAILABLE:
            print("Warning: Windows screen capture protection not available")
            return
        
        try:
            # Get the window handle
            hwnd = int(self.winId())
            
            # Set WDA_EXCLUDEFROMCAPTURE to prevent screen capture
            user32 = ctypes.windll.user32
            result = user32.SetWindowDisplayAffinity(hwnd, WDA_EXCLUDEFROMCAPTURE)
            
            if result:
                print(f"✓ Screen capture protection ENABLED for window {hwnd}")
                # Log the protection activation
                self.api_client.log_security_event(
                    activity_type="screen_capture_protection_enabled",
                    metadata={
                        "document_id": self.document_data.get('id'),
                        "document_name": self.filename,
                        "classification": self.classification,
                        "window_handle": hwnd,
                        "protection_type": "WDA_EXCLUDEFROMCAPTURE"
                    }
                )
            else:
                error_code = ctypes.get_last_error()
                print(f"✗ Failed to enable screen capture protection. Error code: {error_code}")
                # Still log the attempt
                self.api_client.log_security_event(
                    activity_type="screen_capture_protection_failed",
                    metadata={
                        "document_id": self.document_data.get('id'),
                        "document_name": self.filename,
                        "classification": self.classification,
                        "error_code": error_code
                    }
                )
        except Exception as e:
            print(f"Error applying screen capture protection: {str(e)}")
    
    def _get_classification_color(self, classification: str) -> str:
        """Get color based on classification level."""
        colors = {
            'public': '#28a745',
            'internal': '#17a2b8',
            'confidential': '#dc3545',
            'unclassified': '#6c757d'
        }
        return colors.get(classification.lower(), '#6c757d')
    
    def load_document(self):
        """Load document based on file type."""
        try:
            if self.file_extension == '.txt':
                self._load_text_document()
            elif self.file_extension == '.pdf':
                self._load_pdf_document()
            elif self.file_extension in ['.docx', '.doc']:
                self._load_docx_document()
            else:
                self._show_unsupported_format()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load document: {str(e)}")
    
    def _load_text_document(self):
        """Load and display text file."""
        try:
            text_content = self.file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text_content = self.file_content.decode('latin-1')
            except:
                text_content = "Error: Unable to decode text file."
        
        text_edit = QTextEdit()
        text_edit.setPlainText(text_content)
        text_edit.setReadOnly(True)
        text_edit.setStyleSheet("""
            QTextEdit {
                background-color: white;
                font-family: 'Courier New', monospace;
                font-size: 12px;
                padding: 15px;
                border: 1px solid #ddd;
            }
        """)
        text_edit.setContextMenuPolicy(Qt.ContextMenuPolicy.NoContextMenu)
        self.content_layout.addWidget(text_edit)
    
    def _load_pdf_document(self):
        """Load and display PDF file."""
        try:
            import fitz  # PyMuPDF
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(self.file_content)
                tmp_path = tmp_file.name
            
            try:
                pdf_document = fitz.open(tmp_path)
                scroll_area = QScrollArea()
                scroll_area.setWidgetResizable(True)
                scroll_content = QWidget()
                scroll_layout = QVBoxLayout(scroll_content)
                
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    zoom = 2.0
                    mat = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=mat)
                    
                    img = QImage(pix.samples, pix.width, pix.height, 
                                pix.stride, QImage.Format.Format_RGB888)
                    
                    page_label = QLabel()
                    page_label.setPixmap(QPixmap.fromImage(img))
                    page_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    page_label.setStyleSheet("background-color: white; margin: 10px;")
                    scroll_layout.addWidget(page_label)
                    
                    if page_num < len(pdf_document) - 1:
                        separator = QLabel()
                        separator.setFixedHeight(2)
                        separator.setStyleSheet("background-color: #ddd;")
                        scroll_layout.addWidget(separator)
                
                scroll_layout.addStretch()
                scroll_area.setWidget(scroll_content)
                self.content_layout.addWidget(scroll_area)
                pdf_document.close()
            finally:
                try:
                    os.unlink(tmp_path)
                except:
                    pass
        except ImportError:
            self._show_pdf_fallback()
    
    def _show_pdf_fallback(self):
        """Show fallback message for PDF."""
        label = QLabel(f"""
            <div style='text-align: center; padding: 50px;'>
                <h2>📄 PDF Viewer</h2>
                <p><b>Filename:</b> {self.filename}</p>
                <p>Install PyMuPDF: <code>pip install PyMuPDF</code></p>
            </div>
        """)
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.content_layout.addWidget(label)
    
    def _load_docx_document(self):
        """Load and display DOCX file."""
        try:
            from docx import Document as DocxDocument
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(self.file_content)
                tmp_path = tmp_file.name
            
            try:
                doc = DocxDocument(tmp_path)
                text_edit = QTextEdit()
                text_edit.setReadOnly(True)
                text_edit.setContextMenuPolicy(Qt.ContextMenuPolicy.NoContextMenu)
                
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        full_text.append(" | ".join(row_text))
                
                text_edit.setPlainText("\n".join(full_text))
                text_edit.setStyleSheet("""
                    QTextEdit {
                        background-color: white;
                        font-family: Arial, sans-serif;
                        font-size: 12px;
                        padding: 20px;
                        border: 1px solid #ddd;
                    }
                """)
                self.content_layout.addWidget(text_edit)
            finally:
                try:
                    os.unlink(tmp_path)
                except:
                    pass
        except ImportError:
            self._show_docx_fallback()
    
    def _show_docx_fallback(self):
        """Show fallback message for DOCX."""
        label = QLabel(f"""
            <div style='text-align: center; padding: 50px;'>
                <h2>📝 DOCX Viewer</h2>
                <p><b>Filename:</b> {self.filename}</p>
                <p>Install python-docx: <code>pip install python-docx</code></p>
            </div>
        """)
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.content_layout.addWidget(label)
    
    def _show_unsupported_format(self):
        """Show message for unsupported formats."""
        label = QLabel(f"""
            <div style='text-align: center; padding: 50px;'>
                <h2>⚠️ Unsupported Format</h2>
                <p><b>Filename:</b> {self.filename}</p>
                <p>Supported: .txt, .pdf, .docx, .doc</p>
            </div>
        """)
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.content_layout.addWidget(label)
    
    def eventFilter(self, obj, event):
        """Event filter to catch all key events including system shortcuts."""
        if event.type() == event.Type.KeyPress and self.classification.lower() == 'confidential':
            # Check for Win+Shift+S (Snipping Tool) - Qt.Key.Key_Meta is Windows key
            if (event.modifiers() & Qt.KeyboardModifier.ShiftModifier and 
                event.modifiers() & Qt.KeyboardModifier.MetaModifier and 
                event.key() == Qt.Key.Key_S):
                self._handle_screenshot_attempt("Windows Snipping Tool (Win+Shift+S)")
                return True  # Block the event
                
        return super().eventFilter(obj, event)
    
    def _handle_screenshot_attempt(self, method: str):
        """Handle screenshot attempt detection and logging."""
        # Log to backend
        self.api_client.log_security_event(
            activity_type="screenshot_attempt",
            metadata={
                "document_id": self.document_data.get('id'),
                "document_name": self.filename,
                "classification": self.classification,
                "method": method
            }
        )
        
        # Show immediate warning
        msg = QMessageBox(self)
        msg.setIcon(QMessageBox.Icon.Warning)
        msg.setWindowTitle("🚫 Screenshot Blocked")
        msg.setText("⚠️ Screenshot attempts are NOT allowed for confidential documents.")
        msg.setInformativeText(f"Attempted method: {method}\n\nThis security event has been logged and reported.")
        msg.setStandardButtons(QMessageBox.StandardButton.Ok)
        msg.setDefaultButton(QMessageBox.StandardButton.Ok)
        msg.exec()
    
    def keyPressEvent(self, event):
        """Disable copying shortcuts and detect screenshot attempts."""
        # Detect Print Screen key
        if event.key() == Qt.Key.Key_Print:
            if self.classification.lower() == 'confidential':
                self._handle_screenshot_attempt("Print Screen")
            event.ignore()
            return
        
        # Block all Ctrl/Cmd shortcuts to prevent copying
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            event.ignore()
        else:
            super().keyPressEvent(event)
    
    def closeEvent(self, event):
        """Clean up when closing."""
        if self.detection_worker is not None and self.is_monitoring:
            self.detection_worker.stop()
            self.detection_worker.wait()  # Wait for thread to finish
        event.accept()
    
    def resizeEvent(self, event):
        """Ensure overlays resize with window."""
        if hasattr(self, 'security_overlay'):
            self.security_overlay.setGeometry(self.rect())
        if hasattr(self, 'watermark_overlay'):
            self.watermark_overlay.setGeometry(self.rect())
        super().resizeEvent(event)
