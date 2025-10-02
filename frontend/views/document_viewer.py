from PyQt6.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, QPushButton, 
                             QTextEdit, QLabel, QMessageBox, QWidget, QScrollArea)
from PyQt6.QtCore import Qt, QBuffer, QByteArray
from PyQt6.QtGui import QPixmap, QTextDocument, QImage, QPageSize
import qtawesome as qta
import tempfile
import os
from pathlib import Path


class DocumentViewer(QDialog):
    """
    Secure document viewer for PDF, DOCX, and TXT files.
    Documents are displayed in read-only mode and cannot be saved or copied.
    """
    
    def __init__(self, document_data: dict, file_content: bytes, parent=None):
        super().__init__(parent)
        self.document_data = document_data
        self.file_content = file_content
        self.filename = document_data.get('filename', 'Document')
        self.file_extension = Path(self.filename).suffix.lower()
        
        self.setWindowTitle(f"View Document - {self.filename}")
        self.setMinimumSize(900, 700)
        
        # Set window flags to ensure maximize button is available
        self.setWindowFlags(
            Qt.WindowType.Window |
            Qt.WindowType.WindowMaximizeButtonHint |
            Qt.WindowType.WindowMinimizeButtonHint |
            Qt.WindowType.WindowCloseButtonHint
        )
        
        self.setup_ui()
        self.load_document()
        
        # Open the window in maximized state
        self.showMaximized()
    
    def setup_ui(self):
        layout = QVBoxLayout(self)
        
        # Header with document info
        header_layout = QHBoxLayout()
        
        info_label = QLabel(f"<b>Document:</b> {self.filename}")
        info_label.setStyleSheet("font-size: 14px; padding: 10px;")
        header_layout.addWidget(info_label)
        
        classification = self.document_data.get('classification', 'unclassified')
        classification_label = QLabel(f"<b>Classification:</b> {classification.upper()}")
        classification_label.setStyleSheet(f"""
            font-size: 14px; 
            padding: 10px; 
            background-color: {self._get_classification_color(classification)};
            border-radius: 5px;
            color: white;
            font-weight: bold;
        """)
        header_layout.addWidget(classification_label)
        
        header_layout.addStretch()
        
        close_button = QPushButton(qta.icon('fa5s.times', color='red'), "Close")
        close_button.clicked.connect(self.close)
        header_layout.addWidget(close_button)
        
        layout.addLayout(header_layout)
        
        # Content area (will be populated based on file type)
        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        layout.addWidget(self.content_widget)
        
        # Warning label
        warning_label = QLabel("⚠️ This document is displayed in read-only mode for security purposes.")
        warning_label.setStyleSheet("""
            background-color: #fff3cd;
            color: #856404;
            padding: 10px;
            border: 1px solid #ffc107;
            border-radius: 5px;
            font-size: 12px;
        """)
        warning_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(warning_label)
    
    def _get_classification_color(self, classification: str) -> str:
        """Get color based on classification level."""
        colors = {
            'public': '#28a745',
            'internal': '#17a2b8',
            'confidential': '#dc3545',
            'unclassified': '#6c757d'
        }
        return colors.get(classification.lower(), '#6c757d')
    
    def load_document(self):
        """Load document based on file type."""
        try:
            if self.file_extension == '.txt':
                self._load_text_document()
            elif self.file_extension == '.pdf':
                self._load_pdf_document()
            elif self.file_extension in ['.docx', '.doc']:
                self._load_docx_document()
            else:
                self._show_unsupported_format()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load document: {str(e)}")
    
    def _load_text_document(self):
        """Load and display text file."""
        try:
            # Try to decode the content
            text_content = self.file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text_content = self.file_content.decode('latin-1')
            except:
                text_content = "Error: Unable to decode text file."
        
        text_edit = QTextEdit()
        text_edit.setPlainText(text_content)
        text_edit.setReadOnly(True)
        text_edit.setStyleSheet("""
            QTextEdit {
                background-color: white;
                font-family: 'Courier New', monospace;
                font-size: 12px;
                padding: 15px;
                border: 1px solid #ddd;
            }
        """)
        
        # Disable context menu to prevent copying
        text_edit.setContextMenuPolicy(Qt.ContextMenuPolicy.NoContextMenu)
        
        self.content_layout.addWidget(text_edit)
    
    def _load_pdf_document(self):
        """Load and display PDF file."""
        try:
            # Try to use PyMuPDF (fitz) for better PDF rendering
            import fitz  # PyMuPDF
            
            # Create a temporary file for PDF
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(self.file_content)
                tmp_path = tmp_file.name
            
            try:
                # Open PDF with PyMuPDF
                pdf_document = fitz.open(tmp_path)
                
                # Create scroll area for pages
                scroll_area = QScrollArea()
                scroll_area.setWidgetResizable(True)
                scroll_content = QWidget()
                scroll_layout = QVBoxLayout(scroll_content)
                
                # Render each page
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    
                    # Render page to pixmap
                    zoom = 2.0  # Increase resolution
                    mat = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert to QImage
                    img = QImage(pix.samples, pix.width, pix.height, 
                                pix.stride, QImage.Format.Format_RGB888)
                    
                    # Create label to display image
                    page_label = QLabel()
                    page_label.setPixmap(QPixmap.fromImage(img))
                    page_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    page_label.setStyleSheet("background-color: white; margin: 10px;")
                    
                    scroll_layout.addWidget(page_label)
                    
                    # Add page separator
                    if page_num < len(pdf_document) - 1:
                        separator = QLabel()
                        separator.setFixedHeight(2)
                        separator.setStyleSheet("background-color: #ddd;")
                        scroll_layout.addWidget(separator)
                
                scroll_layout.addStretch()
                scroll_area.setWidget(scroll_content)
                self.content_layout.addWidget(scroll_area)
                
                pdf_document.close()
            finally:
                # Clean up temporary file
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                    
        except ImportError:
            # Fallback: Show message if PyMuPDF is not installed
            self._show_pdf_fallback()
    
    def _show_pdf_fallback(self):
        """Show fallback message for PDF if PyMuPDF is not available."""
        label = QLabel()
        label.setText(f"""
            <div style='text-align: center; padding: 50px;'>
                <h2>📄 PDF Viewer</h2>
                <p><b>Filename:</b> {self.filename}</p>
                <p><b>Size:</b> {len(self.file_content)} bytes</p>
                <br>
                <p>To view PDF files, install PyMuPDF:</p>
                <code>pip install PyMuPDF</code>
                <br><br>
                <p style='color: #666;'>PDF viewing requires PyMuPDF library for secure rendering.</p>
            </div>
        """)
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setWordWrap(True)
        self.content_layout.addWidget(label)
    
    def _load_docx_document(self):
        """Load and display DOCX file."""
        try:
            from docx import Document as DocxDocument
            
            # Create a temporary file for DOCX
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(self.file_content)
                tmp_path = tmp_file.name
            
            try:
                # Open DOCX
                doc = DocxDocument(tmp_path)
                
                # Create text edit to display content
                text_edit = QTextEdit()
                text_edit.setReadOnly(True)
                text_edit.setContextMenuPolicy(Qt.ContextMenuPolicy.NoContextMenu)
                
                # Extract text from paragraphs
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text)
                        full_text.append(" | ".join(row_text))
                
                text_edit.setPlainText("\n".join(full_text))
                text_edit.setStyleSheet("""
                    QTextEdit {
                        background-color: white;
                        font-family: Arial, sans-serif;
                        font-size: 12px;
                        padding: 20px;
                        border: 1px solid #ddd;
                        line-height: 1.6;
                    }
                """)
                
                self.content_layout.addWidget(text_edit)
            finally:
                # Clean up temporary file
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                    
        except ImportError:
            self._show_docx_fallback()
    
    def _show_docx_fallback(self):
        """Show fallback message for DOCX if python-docx is not available."""
        label = QLabel()
        label.setText(f"""
            <div style='text-align: center; padding: 50px;'>
                <h2>📝 DOCX Viewer</h2>
                <p><b>Filename:</b> {self.filename}</p>
                <p><b>Size:</b> {len(self.file_content)} bytes</p>
                <br>
                <p>To view DOCX files, install python-docx:</p>
                <code>pip install python-docx</code>
                <br><br>
                <p style='color: #666;'>DOCX viewing requires python-docx library for secure rendering.</p>
            </div>
        """)
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setWordWrap(True)
        self.content_layout.addWidget(label)
    
    def _show_unsupported_format(self):
        """Show message for unsupported file formats."""
        label = QLabel()
        label.setText(f"""
            <div style='text-align: center; padding: 50px;'>
                <h2>⚠️ Unsupported Format</h2>
                <p><b>Filename:</b> {self.filename}</p>
                <p><b>Extension:</b> {self.file_extension}</p>
                <br>
                <p>This file format is not supported for viewing.</p>
                <p>Supported formats: .txt, .pdf, .docx, .doc</p>
            </div>
        """)
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setWordWrap(True)
        self.content_layout.addWidget(label)
    
    def keyPressEvent(self, event):
        """Override key press to disable copying shortcuts."""
        # Disable Ctrl+C, Ctrl+A, Ctrl+S, etc.
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            event.ignore()
        else:
            super().keyPressEvent(event)
